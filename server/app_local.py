import os, json, sqlite3, tempfile, csv, random, logging
from typing import List, Optional
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from server.routes.upload_questions_excel import router as upload_questions_router

logging.basicConfig(level=logging.INFO, format='%(asctime)s [%(levelname)s] %(message)s')


# Use a writable DB path: next to EXE if frozen, else next to this file
def get_db_path():
    if getattr(sys, 'frozen', False):
        # Use LOCALAPPDATA for writable DB location
        local_appdata = os.environ.get('LOCALAPPDATA', os.path.expanduser('~'))
        db_dir = os.path.join(local_appdata, 'IDCS-QP-Generator')
        os.makedirs(db_dir, exist_ok=True)
        return os.path.join(db_dir, 'local_store.db')
    else:
        return os.path.join(os.path.dirname(__file__), 'local_store.db')

import sys
DB_PATH = get_db_path()

def get_conn():
    return sqlite3.connect(DB_PATH)

def init_db():
    conn = get_conn(); cur = conn.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS templates(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            description TEXT,
            total_marks INTEGER,
            instructions TEXT,
            sections TEXT
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS question_bank_titles(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT UNIQUE NOT NULL
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS question_bank(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            question_text TEXT NOT NULL,
            type TEXT NOT NULL,
            options TEXT,
            correct_answer TEXT,
            answer_text TEXT,
            btl INTEGER,
            marks INTEGER,
            status TEXT,
            chapter TEXT,
            course_outcomes TEXT,
            title_id INTEGER,
            FOREIGN KEY(title_id) REFERENCES question_bank_titles(id)
        )
    """)
    conn.commit(); conn.close()

init_db()

app = FastAPI()
app.include_router(upload_questions_router, prefix="/api")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Templates
@app.post('/api/templates')
async def create_template(name: str = Form(...), description: str = Form(""), total_marks: int = Form(100), instructions: str = Form(""), sections: str = Form("[]")):
    try:
        conn = get_conn(); cur = conn.cursor()
        cur.execute("INSERT INTO templates(name,description,total_marks,instructions,sections) VALUES (?,?,?,?,?)", (name,description,total_marks,instructions,sections))
        tid = cur.lastrowid
        conn.commit(); conn.close()
        return {'id': tid}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f'Template insert failed: {e}')

@app.post('/api/templates/update')
async def update_template(id: int = Form(...), name: str = Form(...), description: str = Form(""), total_marks: int = Form(100), instructions: str = Form(""), sections: str = Form("[]")):
    try:
        conn = get_conn(); cur = conn.cursor()
        cur.execute("UPDATE templates SET name=?,description=?,total_marks=?,instructions=?,sections=? WHERE id=?", (name,description,total_marks,instructions,sections,id))
        if cur.rowcount == 0:
            raise HTTPException(status_code=404, detail='Template not found')
        conn.commit(); conn.close()
        return {'updated': id}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f'Template update failed: {e}')

@app.get('/api/templates')
async def list_templates():
    conn = get_conn(); cur = conn.cursor()
    rows = cur.execute('SELECT id,name,description,total_marks,instructions,sections FROM templates ORDER BY id DESC').fetchall()
    conn.close()
    return [{
        'id': r[0], 'name': r[1], 'description': r[2], 'total_marks': r[3], 'instructions': r[4], 'sections': json.loads(r[5] or '[]')
    } for r in rows]

# Question bank titles
@app.post('/api/question-bank-titles')
async def create_title(title: str = Form(...)):
    try:
        conn = get_conn(); cur = conn.cursor()
        cur.execute('INSERT OR IGNORE INTO question_bank_titles(title) VALUES (?)', (title,))
        cur.execute('SELECT id FROM question_bank_titles WHERE title=?', (title,))
        row = cur.fetchone(); conn.commit(); conn.close()
        return {'id': row[0], 'title': title}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f'Title insert failed: {e}')

@app.get('/api/question-bank-titles')
async def list_titles():
    conn = get_conn(); cur = conn.cursor()
    rows = cur.execute('SELECT id,title FROM question_bank_titles ORDER BY title').fetchall(); conn.close()
    return [{'id': r[0], 'title': r[1]} for r in rows]

# Question bank
@app.post('/api/question-bank/bulk')
async def bulk_insert_questions(title_id: int = Form(...), status: str = Form('pending'), payload: str = Form(...)):
    try:
        data = json.loads(payload)
        if not isinstance(data, list):
            raise HTTPException(status_code=400, detail='payload must be a JSON list')
        conn = get_conn(); cur = conn.cursor()
        inserted = 0
        failed = []
        for idx, q in enumerate(data):
            try:
                # Remove images from payload to avoid huge blobs in DB table
                if isinstance(q, dict) and 'images' in q:
                    q.pop('images', None)
                opts = q.get('options') if isinstance(q, dict) else None
                cur.execute("""INSERT INTO question_bank(question_text,type,options,correct_answer,answer_text,btl,marks,status,chapter,course_outcomes,title_id)
                    VALUES (?,?,?,?,?,?,?,?,?,?,?)""",
                    (q.get('question_text','') if isinstance(q, dict) else '', q.get('type','objective') if isinstance(q, dict) else 'descriptive', json.dumps(opts) if opts else None,
                     q.get('correct_answer') if isinstance(q, dict) else None, q.get('answer_text','') if isinstance(q, dict) else None, q.get('btl',2) if isinstance(q, dict) else 2, q.get('marks',1) if isinstance(q, dict) else 1, status,
                     q.get('chapter') if isinstance(q, dict) else None, q.get('course_outcomes') if isinstance(q, dict) else None, title_id))
                inserted += 1
            except Exception as e:
                logging.exception('Failed inserting question index %s: %s', idx, e)
                failed.append({'index': idx, 'error': str(e)})
        conn.commit(); conn.close()
        return {'inserted': inserted, 'failed': failed}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f'Bulk insert failed: {e}')

@app.get('/api/question-bank')
async def list_questions(status: Optional[str] = None, title_id: Optional[int] = None, title: Optional[str] = None):
    conn = get_conn(); cur = conn.cursor()
    base = 'SELECT id,question_text,type,options,correct_answer,answer_text,btl,marks,status,chapter,course_outcomes,title_id FROM question_bank WHERE 1=1'
    params: List = []
    if status:
        base += ' AND status=?'; params.append(status)
    if title_id:
        base += ' AND title_id=?'; params.append(title_id)
    if title:
        base += ' AND title=?'; params.append(title)
    rows = cur.execute(base, params).fetchall(); conn.close()
    return [{
        'id': r[0], 'question_text': r[1], 'type': r[2], 'options': json.loads(r[3]) if r[3] else None,
        'correct_answer': r[4], 'answer_text': r[5], 'btl': r[6], 'marks': r[7], 'status': r[8],
        'chapter': r[9], 'course_outcomes': r[10], 'title_id': r[11]
    } for r in rows]

@app.post('/api/question-bank/update-status')
async def update_question_status(ids: str = Form(...), status: str = Form(...)):
    try:
        id_list = [int(x) for x in ids.split(',') if x.strip().isdigit()]
        if not id_list:
            raise HTTPException(status_code=400, detail='No valid ids')
        conn = get_conn(); cur = conn.cursor()
        cur.executemany('UPDATE question_bank SET status=? WHERE id=?', [(status, i) for i in id_list])
        conn.commit(); conn.close()
        return {'updated': len(id_list), 'status': status}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f'Status update failed: {e}')


# Admin-only: seed sample pending questions for a given title_id
@app.post('/api/admin/seed-question-bank')
async def seed_question_bank(title_id: str = Form(...), count: int = Form(3), admin_secret: str = Form(...)):
    """Seed `count` pending questions with the given `title_id`.
    This endpoint is protected by an admin secret (set ADMIN_SECRET env var).
    Use only for testing/local development.
    """
    try:
        expected = os.environ.get('ADMIN_SECRET', 'dev-secret')
        if admin_secret != expected:
            raise HTTPException(status_code=403, detail='Invalid admin secret')
        conn = get_conn(); cur = conn.cursor()
        inserted_ids = []
        for i in range(int(count)):
            qtext = f'SEED: {title_id} sample {i+1}'
            cur.execute("""INSERT INTO question_bank(question_text,type,options,correct_answer,answer_text,btl,marks,status,chapter,course_outcomes,title_id)
                VALUES (?,?,?,?,?,?,?,?,?,?,?)""",
                        (qtext, 'objective', None, None, None, 2, 1, 'pending', None, None, title_id))
            inserted_ids.append(cur.lastrowid)
        conn.commit(); conn.close()
        return {'inserted': len(inserted_ids), 'ids': inserted_ids}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f'Seed failed: {e}')

# Template file upload/scan
@app.post('/api/template/upload')
async def upload_template(file: UploadFile = File(...)):
    ext = os.path.splitext(file.filename)[1].lower(); content_lines: List[str] = []
    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        tmp.write(await file.read()); tmp_path = tmp.name
    try:
        if ext == '.txt':
            with open(tmp_path, encoding='utf-8') as f:
                content_lines = [ln.strip() for ln in f if ln.strip()]
        elif ext == '.csv':
            with open(tmp_path, encoding='utf-8') as f:
                for row in csv.reader(f):
                    content_lines.append(', '.join(row))
        elif ext == '.docx':
            from docx import Document
            doc = Document(tmp_path)
            for para in doc.paragraphs:
                t = para.text.strip()
                if t:
                    content_lines.append(t)
        else:
            return JSONResponse(status_code=400, content={'error':'Unsupported file type'})
    finally:
        os.remove(tmp_path)
    return {'lines': content_lines}

@app.post('/api/template/scan-docx')
async def scan_docx(file: UploadFile = File(...)):
    from docx import Document
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
        tmp.write(await file.read()); tmp_path = tmp.name
    questions = []
    try:
        doc = Document(tmp_path); part=None
        import base64
        # Build a map of image related parts (rid -> data-uri)
        image_map = {}
        try:
            for rid, part_obj in doc.part.related_parts.items():
                ct = getattr(part_obj, 'content_type', '') or ''
                if ct.startswith('image/'):
                    blob = getattr(part_obj, 'blob', None)
                    if blob:
                        image_map[rid] = f"data:{ct};base64,{base64.b64encode(blob).decode('ascii')}"
        except Exception:
            image_map = {}

        def extract_images_from_paragraph(p):
            imgs = []
            try:
                # Brute force search for any attribute value that matches a known image RID
                for elem in p._element.iter():
                    for attrib_name, attrib_value in elem.attrib.items():
                        if attrib_value in image_map:
                            if image_map[attrib_value] not in imgs:
                                imgs.append(image_map[attrib_value])
                                if logger:
                                    logger.info('[scan-docx] found image via attribute %s="%s" on tag %s', attrib_name, attrib_value, elem.tag)
            except Exception as e:
                if logger:
                    logger.error('[scan-docx] error extracting images from paragraph: %s', e)
            return imgs

        def extract_images_from_cell(cell):
            imgs = []
            try:
                # Search paragraphs inside the cell
                for p in cell.paragraphs:
                    imgs.extend(extract_images_from_paragraph(p))
                # also search for r:embed attributes in the cell XML
                ns = {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main', 'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture', 'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'}
                try:
                    embeds = cell._element.xpath('.//@r:embed', namespaces=ns)
                    for rel in embeds:
                        if rel in image_map and image_map[rel] not in imgs:
                            imgs.append(image_map[rel])
                            if logger:
                                logger.info('[scan-docx] cell attached image via r:embed rel=%s', rel)
                except Exception:
                    pass
            except Exception:
                pass
            return imgs
        # Diagnostics: table shapes and paragraph snippets
        try:
            import logging
            logger = logging.getLogger('app_local')
        except Exception:
            logger = None
        table_shapes = []
        sample_table_texts = []
        for t in doc.tables:
            try:
                table_shapes.append(len(t.columns))
                # capture first row text of each table for quick inspection
                rows_txt = []
                for r in t.rows[:3]:
                    cells = [c.text.strip().replace('\n',' ')[:120] for c in r.cells]
                    rows_txt.append(' | '.join(cells))
                sample_table_texts.append(' || '.join(rows_txt))
            except Exception:
                sample_table_texts.append('error')
        para_snippets = [p.text.strip().replace('\n',' ')[:240] for p in doc.paragraphs[:20]]
        # First, try to extract questions from tables (expected template format)
        for table in doc.tables:
            cols = len(table.columns)
            # Part A - 4 or 5 column table (some generators produce 5 cols: Q.No, Question, CO, BTL, Marks)
            if cols == 4 or cols == 5:
                part = 'A'
                for row in table.rows[1:]:
                    # question cell at index 1
                    qcell = row.cells[1].text.strip() if len(row.cells) > 1 else ''
                    co = row.cells[2].text.strip() if len(row.cells) > 2 else ''
                    # BTL and marks shift when 5 columns
                    btl = row.cells[3].text.strip() if len(row.cells) > 3 else ''
                    marks = row.cells[4].text.strip() if cols == 5 and len(row.cells) > 4 else (row.cells[3].text.strip() if len(row.cells) > 3 else '')
                    if qcell:
                        imgs = extract_images_from_cell(row.cells[1]) if len(row.cells) > 1 else []
                        qobj = {'number': len(questions) + 1, 'text': qcell, 'co': co, 'btl': btl, 'marks': marks or 2, 'part': 'A'}
                        if imgs:
                            qobj['images'] = imgs
                        questions.append(qobj)
            # Part B: tables with OR structure - accept 5+ columns as well (some generators use 5 columns)
            elif cols >= 5:
                part = 'B'; rows = table.rows; i = 0
                while i < len(rows):
                    if 'OR' in rows[i].cells[0].text.upper():
                        i += 1; continue
                    if i + 2 < len(rows) and 'OR' in rows[i+1].cells[0].text.upper():
                        # pick question text and metadata from appropriate cells
                        q_a = rows[i].cells[1].text.strip() if len(rows[i].cells) > 1 else ''
                        q_b = rows[i+2].cells[1].text.strip() if len(rows[i+2].cells) > 1 else ''
                        co_a = rows[i].cells[2].text.strip() if len(rows[i].cells) > 2 else ''
                        co_b = rows[i+2].cells[2].text.strip() if len(rows[i+2].cells) > 2 else ''
                        btl_a = rows[i].cells[3].text.strip() if len(rows[i].cells) > 3 else ''
                        btl_b = rows[i+2].cells[3].text.strip() if len(rows[i+2].cells) > 3 else ''
                        marks_a = rows[i].cells[-1].text.strip() if len(rows[i].cells) > 0 else '16'
                        marks_b = rows[i+2].cells[-1].text.strip() if len(rows[i+2].cells) > 0 else '16'
                        base_num = 10 + (len([q for q in questions if q.get('part') == 'B']) // 2) + 1
                        if q_a:
                            imgs_a = extract_images_from_cell(rows[i].cells[1]) if len(rows[i].cells) > 1 else []
                            qobj_a = {'number': f'{base_num}a', 'text': q_a, 'co': co_a, 'btl': btl_a, 'marks': marks_a, 'part': 'B', 'or': False}
                            if imgs_a: qobj_a['images'] = imgs_a
                            questions.append(qobj_a)
                        if q_b:
                            imgs_b = extract_images_from_cell(rows[i+2].cells[1]) if len(rows[i+2].cells) > 1 else []
                            qobj_b = {'number': f'{base_num}b', 'text': q_b, 'co': co_b, 'btl': btl_b, 'marks': marks_b, 'part': 'B', 'or': True}
                            if imgs_b: qobj_b['images'] = imgs_b
                            questions.append(qobj_b)
                        i += 3
                    else:
                        i += 1
        # If no questions found in tables, try a fallback: parse paragraphs for numbered lines
        if not questions:
            import re
            orig_paras = list(doc.paragraphs)
            paras = [p.text.rstrip() for p in orig_paras]
            q_re = re.compile(r"^\s*(\d+[a-zA-Z0-9]*)[\.|\)|\-|:]?\s+(.*)")
            opt_re = re.compile(r"^\s*(?:\(?[A-Za-z0-9]{1,2}\)?[\.|\)]\s*)(.*)")
            ans_re = re.compile(r"^(?:Answer|Ans|Correct|Solution)\s*[:\-]\s*(.*)", re.I)
            i = 0
            while i < len(paras):
                line = paras[i].strip()
                if not line:
                    i += 1; continue
                m = q_re.match(line)
                if m:
                    num = m.group(1)
                    text = m.group(2).strip()
                    options = []
                    answer_text = None
                    j = i + 1
                    # gather following option/answer lines until next question
                    while j < len(paras):
                        nxt = paras[j].strip()
                        if not nxt:
                            j += 1; continue
                        if q_re.match(nxt):
                            break
                        am = ans_re.match(nxt)
                        om = opt_re.match(nxt)
                        if am:
                            answer_text = am.group(1).strip()
                        elif om:
                            options.append(om.group(1).strip())
                        else:
                            # continuation: attach to question text or last option
                            if options:
                                options[-1] = options[-1] + ' ' + nxt
                            else:
                                text = text + ' ' + nxt
                        # also extract images present in this paragraph
                        try:
                            imgs_here = extract_images_from_paragraph(orig_paras[j])
                            if imgs_here:
                                if 'images' not in locals():
                                    images = []
                                images.extend(imgs_here)
                        except Exception:
                            pass
                        j += 1
                    qobj = {'number': num, 'text': text, 'co': None, 'btl': None, 'marks': None, 'part': None}
                    if options:
                        qobj['options'] = options
                        qobj['type'] = 'objective'
                    if answer_text:
                        qobj['answer_text'] = answer_text
                    # attach any collected images
                    if 'images' in locals() and images:
                        qobj['images'] = images
                        del images
                    questions.append(qobj)
                    i = j
                else:
                    i += 1
    finally:
        os.remove(tmp_path)
    diagnostic = {
        'table_count': len(table_shapes),
        'table_shapes': table_shapes,
        'sample_table_texts': sample_table_texts,
        'paragraph_snippets': para_snippets,
        'parsed_questions': len(questions),
        'image_count': sum(len(q.get('images', [])) for q in questions)
    }
    # Summarize related parts (images, media) present in the docx package for debugging
    try:
        related = []
        for rid, part_obj in doc.part.related_parts.items():
            ct = getattr(part_obj, 'content_type', None) or str(type(part_obj))
            blob = getattr(part_obj, 'blob', None)
            size = len(blob) if blob else None
            related.append({'rel': rid, 'content_type': ct, 'size': size})
        diagnostic['related_parts'] = related
    except Exception:
        diagnostic['related_parts'] = []
    # Search all package parts for references to related part ids and for image-like tags
    try:
        references = {rid: [] for rid, _ in doc.part.related_parts.items()}
        image_hint_parts = set()
        pkg = getattr(doc.part, 'package', None)
        if pkg is not None:
            for p in pkg.parts:
                try:
                    blob = getattr(p, 'blob', None)
                    if not blob:
                        continue
                    text = blob.decode('utf-8', errors='ignore')
                    # check for direct references to relationship ids (rIdXXX)
                    for rid in list(references.keys()):
                        if rid in text:
                            references[rid].append(str(getattr(p, 'partname', p.partname)))
                    # check for image-related XML tags
                    if any(k in text for k in ('<a:blip', 'pic:blipFill', '<v:imagedata', 'img src="')):
                        image_hint_parts.add(str(getattr(p, 'partname', p.partname)))
                except Exception:
                    continue
        diagnostic['references'] = references
        diagnostic['image_hint_parts'] = list(image_hint_parts)
    except Exception:
        diagnostic['references'] = {}
        diagnostic['image_hint_parts'] = []
    if logger:
        try:
            logger.info('[scan-docx] diagnostic: %s', diagnostic)
        except Exception:
            pass
    # Return diagnostic for debugging; front-end can ignore if not used
    return {'questions': questions, 'diagnostic': diagnostic}

@app.post('/api/template/generate-docx')
async def generate_docx(
    questions: str = Form(...), dept: str = Form(""), cc: str = Form(""), cn: str = Form(""), qpcode: str = Form(""),
    exam_title: str = Form("B.E., /B.Tech., DEGREE EXAMINATIONS, APRIL/MAY2024"), regulation: str = Form("Regulation 2024"),
    semester: str = Form("Second Semester")
):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()

    def add_bold_line(text: str, center=True, size=12, underline=False, border=False):
        p = doc.add_paragraph(); run = p.add_run(text); run.bold=True; run.font.size=Pt(size); run.underline=underline
        if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if border:
            tbl = doc.add_table(rows=1, cols=1); c=tbl.rows[0].cells[0]; c.text=text
            for para in c.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in para.runs: r.bold=True; r.font.size=Pt(size)
            return tbl
        return p

    def add_line(text: str, center=True, size=12, italic=False):
        p=doc.add_paragraph(); run=p.add_run(text); run.italic=italic; run.font.size=Pt(size)
        if center: p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        return p

    BOX_COUNT=16
    reg_table = doc.add_table(rows=2, cols=BOX_COUNT+1); reg_table.alignment=WD_TABLE_ALIGNMENT.CENTER; reg_table.autofit=False
    def set_table_borders(tbl):
        tbl_xml=tbl._tbl; tblPr=tbl_xml.tblPr
        if tblPr is None:
            from docx.oxml import OxmlElement
            tblPr=OxmlElement('w:tblPr'); tbl_xml.insert(0,tblPr)
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        borders=tblPr.find(qn('w:tblBorders'))
        if borders is None:
            borders=OxmlElement('w:tblBorders'); tblPr.append(borders)
        for side in ('top','left','bottom','right','insideH','insideV'):
            tag=qn(f'w:{side}'); element=borders.find(tag)
            if element is None:
                element=OxmlElement(f'w:{side}'); borders.append(element)
            element.set(qn('w:val'),'single'); element.set(qn('w:sz'),'12'); element.set(qn('w:space'),'0'); element.set(qn('w:color'),'000000')
    set_table_borders(reg_table)

    from docx.shared import Inches
    box_width=Inches(0.25); box_height_twips=360
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    def set_row_height(row, twips):
        tr=row._tr; trPr=tr.get_or_add_trPr()
        for child in trPr.findall(qn('w:trHeight')): trPr.remove(child)
        trHeight=OxmlElement('w:trHeight'); trHeight.set(qn('w:val'), str(twips)); trHeight.set(qn('w:hRule'),'exact'); trPr.append(trHeight)

    label_top=reg_table.rows[0].cells[0]; label_bottom=reg_table.rows[1].cells[0]; label_top.merge(label_bottom)
    label_top.text='Reg.\nNo.:'
    for p in label_top.paragraphs:
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs: r.bold=True; r.font.size=Pt(10)
    for col in range(1,BOX_COUNT+1):
        top=reg_table.rows[0].cells[col]; bottom=reg_table.rows[1].cells[col]; merged=top.merge(bottom); merged.text=''
        merged.width=box_width
        for p in merged.paragraphs: p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    set_row_height(reg_table.rows[0], box_height_twips); set_row_height(reg_table.rows[1], box_height_twips)

    add_bold_line('\nK. RAMAKRISHNAN COLLEGE OF ENGINEERING', True, 14)
    add_bold_line('(AUTONOMOUS)', True, 12)
    add_bold_line(f'Question Paper Code: {qpcode}', True, 12, border=True)
    if exam_title.strip() and not exam_title.lower().startswith('question paper code:'):
        add_bold_line(exam_title, True, 12)
    add_line(semester, True, 11, italic=True)
    if dept: add_line(dept, True, 11, italic=True)
    add_bold_line(f'{cc} – {cn}', True, 12)
    add_line(f'({regulation})', True, 11)

    tm = doc.add_table(rows=1, cols=2); tm.alignment=WD_TABLE_ALIGNMENT.CENTER; tm.autofit=True
    tm.rows[0].cells[0].text='Time: Three Hours'; tm.rows[0].cells[1].text='Maximum Marks: 100 Marks'
    for ci, cell in enumerate(tm.rows[0].cells):
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci==0 else WD_ALIGN_PARAGRAPH.RIGHT
            for r in p.runs: r.font.size=Pt(11)

    doc.add_paragraph('')
    add_bold_line('PART- A                                                                (10 x 2 = 20 Marks)', True, 12)
    table_a = doc.add_table(rows=1, cols=5); table_a.alignment=WD_TABLE_ALIGNMENT.CENTER; table_a.autofit=False
    hdr=table_a.rows[0].cells; hdr[0].text='Q.No.'; hdr[1].text='Answer ALL Questions'; hdr[2].text='CO'; hdr[3].text='BTL'; hdr[4].text='Marks'
    widths_a=[Inches(0.7), Inches(4.2), Inches(0.8), Inches(0.8), Inches(0.8)]
    for i,w in enumerate(widths_a):
        for row in table_a.rows: row.cells[i].width=w
    for c in hdr:
        for p in c.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold=True

    def normalize(raw):
        out=[]
        def push(item):
            if isinstance(item,str):
                try: item=json.loads(item)
                except Exception: item={'text':item}
            if isinstance(item,list):
                for sub in item: push(sub)
            elif isinstance(item,dict): out.append(item)
        push(raw)
        return out if out else []

    try:
        parsed = normalize(questions)
    except Exception:
        parsed=[]

    def first(d:dict, keys:List[str]):
        for k in keys:
            if k in d and str(d[k]).strip(): return str(d[k])
        return ''

    shared_btl=random.choice([3,4,5])
    import re
    for idx,q in enumerate(parsed[:10], start=1):
        cells=table_a.add_row().cells
        for j,w in enumerate(widths_a): cells[j].width=w
        cells[0].text=str(idx)
        txt=first(q,['text','question_text'])
        if txt: txt=re.sub(r'^\s*[DO]\.[\s-]*','',txt,flags=re.IGNORECASE)
        cells[1].text=txt
        co_val=q.get('co') or f'CO{(idx+1)//2}'
        cells[2].text=str(co_val)
        btl_val=q.get('btl') or (shared_btl if idx>4 else random.choice([1,2,3,4,5]))
        cells[3].text=f'BTL{btl_val}'
        cells[4].text='2'
        for p in cells[0].paragraphs+cells[2].paragraphs+cells[3].paragraphs+cells[4].paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER

    add_bold_line('PART – B                          (5 x 16 = 80 Marks)', True, 12)
    table_b=doc.add_table(rows=1, cols=5); table_b.alignment=WD_TABLE_ALIGNMENT.CENTER; table_b.autofit=False
    bh=table_b.rows[0].cells; bh[0].text='Q.No.'; bh[1].text='Question'; bh[2].text='CO'; bh[3].text='BTL'; bh[4].text='Marks'
    widths_b=[Inches(0.9), Inches(4.5), Inches(0.9), Inches(0.9), Inches(1.0)]
    for i,w in enumerate(widths_b):
        for row in table_b.rows: row.cells[i].width=w
    for c in bh:
        for p in c.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold=True

    from collections import defaultdict
    groups=defaultdict(list)
    for q in parsed:
        if str(q.get('part','')).upper()=='B':
            groups[q.get('number')] .append(q)
    def sort_key(k):
        try: return int(str(k).split('.')[0])
        except: return 9999
    for base in sorted(groups.keys(), key=sort_key):
        group=groups[base]; group.sort(key=lambda x: str(x.get('sub','a')))
        for idx_in, q in enumerate(group):
            cells=table_b.add_row().cells
            for i,w in enumerate(widths_b): cells[i].width=w
            sub=q.get('sub'); disp=f'{base}.{sub}' if sub else str(base)
            cells[0].text=disp
            cells[1].text=first(q,['text','question_text'])
            co_val=first(q,['co','course_outcomes']); cells[2].text=co_val
            btl_val=first(q,['btl'])
            if btl_val and not btl_val.upper().startswith('BTL'): btl_val=f'BTL{btl_val}'
            cells[3].text=btl_val
            cells[4].text=first(q,['marks']) or '16'
            for p in cells[0].paragraphs+cells[2].paragraphs+cells[3].paragraphs+cells[4].paragraphs:
                p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            if idx_in==0 and any(str(x.get('sub','')).lower()=='b' for x in group):
                or_row=table_b.add_row().cells
                for i,w in enumerate(widths_b): or_row[i].width=w
                merged=or_row[0]
                for c_idx in range(1,len(or_row)): merged=merged.merge(or_row[c_idx])
                merged.text='(OR)'
                for p in merged.paragraphs:
                    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs: r.bold=True

    # PART-C (optional) - typically single pair 16.a / 16.b with OR
    c_items = [q for q in parsed if str(q.get('part','')).upper()=='C']
    if c_items:
        # Determine counts and projection marks for header
        from collections import defaultdict
        c_groups=defaultdict(list)
        for q in c_items:
            key=str(q.get('number') or q.get('baseNumber') or 16)
            # normalize to base number such as 16
            try:
                key = str(int(str(key).split('.')[0]))
            except Exception:
                key = '16'
            c_groups[key].append(q)
        # Use first question's marks as projection marks
        first_c = c_items[0]
        proj_marks = first_c.get('marks') or '10'
        try:
            count_pairs = len(c_groups.keys())
            total_marks = int(str(proj_marks)) * count_pairs
            add_bold_line(f'PART – C                          ({count_pairs} x {proj_marks} = {total_marks} Marks)', True, 12)
        except Exception:
            add_bold_line('PART – C', True, 12)
        table_c=doc.add_table(rows=1, cols=5); table_c.alignment=WD_TABLE_ALIGNMENT.CENTER; table_c.autofit=False
        ch=table_c.rows[0].cells; ch[0].text='Q.No.'; ch[1].text='Question'; ch[2].text='CO'; ch[3].text='BTL'; ch[4].text='Marks'
        widths_c=[Inches(0.9), Inches(4.5), Inches(0.9), Inches(0.9), Inches(1.0)]
        for i,w in enumerate(widths_c):
            for row in table_c.rows: row.cells[i].width=w
        for c in ch:
            for p in c.paragraphs:
                p.alignment=WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.bold=True
        for base in sorted(c_groups.keys(), key=lambda k: int(k)):
            group=c_groups[base]
            # try to order a then b
            try:
                group.sort(key=lambda q: str(q.get('sub','a')))
            except Exception: pass
            # (a)
            row_a=table_c.add_row().cells
            for i,w in enumerate(widths_c): row_a[i].width=w
            row_a[0].text=f'{base}.a'
            row_a[1].text=first(group[0] if group else {},['text','question_text']) if group else ''
            row_a[2].text=first(group[0] if group else {},['co','course_outcomes']) if group else ''
            btl_val=first(group[0] if group else {},['btl'])
            if btl_val and not btl_val.upper().startswith('BTL'): btl_val=f'BTL{btl_val}'
            row_a[3].text=btl_val
            row_a[4].text=str((group[0] or {}).get('marks','')) if group else ''
            for p in row_a[0].paragraphs + row_a[2].paragraphs + row_a[3].paragraphs + row_a[4].paragraphs:
                p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            # OR row
            or_row=table_c.add_row().cells
            for i,w in enumerate(widths_c): or_row[i].width=w
            merged=or_row[0]
            for ci in range(1,len(or_row)): merged=merged.merge(or_row[ci])
            merged.text='(OR)'
            for p in merged.paragraphs:
                p.alignment=WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.bold=True
            # (b)
            row_b=table_c.add_row().cells
            for i,w in enumerate(widths_c): row_b[i].width=w
            row_b[0].text=f'{base}.b'
            sec = group[1] if len(group)>1 else {}
            row_b[1].text=first(sec,['text','question_text'])
            row_b[2].text=first(sec,['co','course_outcomes'])
            btl_val=first(sec,['btl'])
            if btl_val and not btl_val.upper().startswith('BTL'): btl_val=f'BTL{btl_val}'
            row_b[3].text=btl_val
            row_b[4].text=str(sec.get('marks',''))
            for p in row_b[0].paragraphs + row_b[2].paragraphs + row_b[3].paragraphs + row_b[4].paragraphs:
                p.alignment=WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(' ')
    doc.add_paragraph('******************').bold=True
    doc.add_paragraph(f'  {qpcode}').bold=True
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
        doc.save(tmp.name); path=tmp.name
    return FileResponse(path, filename='question_paper.docx')

if __name__ == '__main__':
    import uvicorn
    uvicorn.run(app, host='0.0.0.0', port=4001)
# For deployment, use: https://idcs-main-kucq.onrender.com
