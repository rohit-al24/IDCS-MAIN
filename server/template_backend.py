import logging

# Setup logging to file and console
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    handlers=[
        logging.FileHandler("template_backend.log", mode="a", encoding="utf-8"),
        logging.StreamHandler()
    ]
)
from logging.handlers import RotatingFileHandler
logger = logging.getLogger("template_backend")
logger.setLevel(logging.INFO)

# Add file handler to log to 'server.log' with rotation
file_handler = RotatingFileHandler("server.log", maxBytes=2*1024*1024, backupCount=3)
file_handler.setLevel(logging.INFO)
formatter = logging.Formatter('%(asctime)s %(levelname)s %(name)s %(message)s')
file_handler.setFormatter(formatter)
if not any(isinstance(h, RotatingFileHandler) for h in logger.handlers):
    logger.addHandler(file_handler)
from fastapi import FastAPI, File, UploadFile, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from typing import List, Optional
import uvicorn
import os
import tempfile
import csv
from docx import Document
from server.routes.upload_questions_excel import router as upload_questions_router

app = FastAPI()

# Register the upload questions router
app.include_router(upload_questions_router, prefix="/api")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.post("/api/template/upload")
async def upload_template(file: UploadFile = File(...)):
    ext = os.path.splitext(file.filename)[1].lower()
    content_lines: List[str] = []
    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        tmp.write(await file.read())
        tmp_path = tmp.name
    try:
        if ext == ".txt":
            with open(tmp_path, encoding="utf-8") as f:
                content_lines = [line.strip() for line in f if line.strip()]
        elif ext == ".csv":
            with open(tmp_path, encoding="utf-8") as f:
                reader = csv.reader(f)
                for row in reader:
                    content_lines.append(", ".join(row))
        elif ext == ".docx":
            doc = Document(tmp_path)
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    content_lines.append(text)
        else:
            return JSONResponse(status_code=400, content={"error": "Unsupported file type"})
    finally:
        os.remove(tmp_path)
    return {"lines": content_lines}

@app.post("/api/template/scan-docx")
async def scan_docx(file: UploadFile = File(...)):
    import re
    from docx import Document
    import tempfile
    import os
    questions = []
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
        tmp.write(await file.read())
        tmp_path = tmp.name
    try:
        doc = Document(tmp_path)
        part = None
        import base64
        # Build a map of image related parts (rid -> data-uri)
        image_map = {}
        try:
            for rid, part_obj in doc.part.related_parts.items():
                ct = getattr(part_obj, 'content_type', '') or ''
                if ct.startswith('image/'):
                    blob = getattr(part_obj, 'blob', None)
                    if blob:
                        image_map[rid] = f"data:{ct};base64,{base64.b64encode(blob).decode('ascii')}"
        except Exception:
            image_map = {}

        def extract_images_from_paragraph(p):
            imgs = []
            try:
                # Brute force search for any attribute value that matches a known image RID
                for elem in p._element.iter():
                    for attrib_name, attrib_value in elem.attrib.items():
                        if attrib_value in image_map:
                            if image_map[attrib_value] not in imgs:
                                imgs.append(image_map[attrib_value])
                                logger.info('template_backend: found image via attribute %s="%s" on tag %s', attrib_name, attrib_value, elem.tag)
            except Exception as e:
                logger.error('template_backend: error extracting images from paragraph: %s', e)
            return imgs

        def extract_images_from_cell(cell):
            imgs = []
            try:
                for p in cell.paragraphs:
                    imgs.extend(extract_images_from_paragraph(p))
                ns = {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main', 'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture', 'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'}
                try:
                    embeds = cell._element.xpath('.//@r:embed', namespaces=ns)
                    for rel in embeds:
                        if rel in image_map and image_map[rel] not in imgs:
                            imgs.append(image_map[rel])
                            logger.info('template_backend: cell attached image via r:embed rel=%s', rel)
                except Exception:
                    pass
            except Exception:
                pass
            return imgs
        # Diagnostics capture for debugging
        try:
            table_shapes = []
            sample_table_texts = []
            for t in doc.tables:
                table_shapes.append(len(t.columns))
                rows_txt = []
                for r in t.rows[:3]:
                    cells = [c.text.strip().replace('\n',' ')[:120] for c in r.cells]
                    rows_txt.append(' | '.join(cells))
                sample_table_texts.append(' || '.join(rows_txt))
            para_snippets = [p.text.strip().replace('\n',' ')[:240] for p in doc.paragraphs[:20]]
        except Exception:
            table_shapes = []
            sample_table_texts = []
            para_snippets = []
        for table in doc.tables:
            # PART-A: 4 or 5 column table (some templates use 5 columns: Q.No, Question, CO, BTL, Marks)
            cols = len(table.columns)
            if cols == 4 or cols == 5:
                part = 'A'
                for row in table.rows[1:]:
                    qtext = row.cells[1].text.strip() if len(row.cells) > 1 else ''
                    co = row.cells[2].text.strip() if len(row.cells) > 2 else ''
                    btl = row.cells[3].text.strip() if len(row.cells) > 3 else ''
                    marks = row.cells[4].text.strip() if cols == 5 and len(row.cells) > 4 else (row.cells[3].text.strip() if len(row.cells) > 3 else '')
                    number = len(questions) + 1
                    if qtext:
                        imgs = extract_images_from_cell(row.cells[1]) if len(row.cells) > 1 else []
                        qobj = {
                            'number': number,
                            'text': qtext,
                            'co': co,
                            'btl': btl,
                            'marks': marks or 2,
                            'part': part
                        }
                        if imgs: qobj['images'] = imgs
                        questions.append(qobj)
            # PART-B: tables with OR structure. Accept 5+ columns (some generators use 5 columns)
            elif cols >= 5:
                part = 'B'
                rows = table.rows
                i = 0
                while i < len(rows):
                    row = rows[i]
                    if 'OR' in row.cells[0].text.upper():
                        i += 1
                        continue
                    if i+2 < len(rows) and 'OR' in rows[i+1].cells[0].text.upper():
                        # a/b pair
                        q_a = row.cells[1].text.strip() if len(row.cells) > 1 else ''
                        co_a = row.cells[2].text.strip() if len(row.cells) > 2 else ''
                        btl_a = row.cells[3].text.strip() if len(row.cells) > 3 else ''
                        marks_a = row.cells[-1].text.strip() if len(row.cells) > 0 else '16'
                        q_b = rows[i+2].cells[1].text.strip() if len(rows[i+2].cells) > 1 else ''
                        co_b = rows[i+2].cells[2].text.strip() if len(rows[i+2].cells) > 2 else ''
                        btl_b = rows[i+2].cells[3].text.strip() if len(rows[i+2].cells) > 3 else ''
                        marks_b = rows[i+2].cells[-1].text.strip() if len(rows[i+2].cells) > 0 else '16'
                        number = 10 + (len(questions) // 2) + 1
                        if q_a:
                            imgs_a = extract_images_from_cell(row.cells[1]) if len(row.cells) > 1 else []
                            qobj_a = {
                                'number': f'{number}a',
                                'text': q_a,
                                'co': co_a,
                                'btl': btl_a,
                                'marks': marks_a,
                                'part': part,
                                'or': False
                            }
                            if imgs_a: qobj_a['images'] = imgs_a
                            questions.append(qobj_a)
                        if q_b:
                            imgs_b = extract_images_from_cell(rows[i+2].cells[1]) if len(rows[i+2].cells) > 1 else []
                            qobj_b = {
                                'number': f'{number}b',
                                'text': q_b,
                                'co': co_b,
                                'btl': btl_b,
                                'marks': marks_b,
                                'part': part,
                                'or': True
                            }
                            if imgs_b: qobj_b['images'] = imgs_b
                            questions.append(qobj_b)
                        i += 3
                    else:
                        i += 1
    finally:
        os.remove(tmp_path)
    # Fallback: if no questions found from tables, try paragraph-based parsing
    if not questions:
        import re
        orig_paras = list(doc.paragraphs)
        paras = [p.text.rstrip() for p in orig_paras]
        q_re = re.compile(r"^\s*(\d+[a-zA-Z0-9]*)[\.|\)|\-|:]?\s+(.*)")
        opt_re = re.compile(r"^\s*(?:\(?[A-Za-z0-9]{1,2}\)?[\.|\)]\s*)(.*)")
        ans_re = re.compile(r"^(?:Answer|Ans|Correct|Solution)\s*[:\-]\s*(.*)", re.I)
        i = 0
        while i < len(paras):
            line = paras[i].strip()
            if not line:
                i += 1; continue
            m = q_re.match(line)
            if m:
                num = m.group(1)
                text = m.group(2).strip()
                options = []
                answer_text = None
                j = i + 1
                while j < len(paras):
                    nxt = paras[j].strip()
                    if not nxt:
                        j += 1; continue
                    if q_re.match(nxt):
                        break
                    am = ans_re.match(nxt)
                    om = opt_re.match(nxt)
                    if am:
                        answer_text = am.group(1).strip()
                    elif om:
                        options.append(om.group(1).strip())
                    else:
                        if options:
                            options[-1] = options[-1] + ' ' + nxt
                        else:
                            text = text + ' ' + nxt
                    # also extract images present in this paragraph
                    try:
                        imgs_here = extract_images_from_paragraph(orig_paras[j])
                        if imgs_here:
                            if 'images' not in locals():
                                images = []
                            images.extend(imgs_here)
                    except Exception:
                        pass
                    j += 1
                qobj = {'number': num, 'text': text, 'co': None, 'btl': None, 'marks': None, 'part': None}
                if options:
                    qobj['options'] = options
                    qobj['type'] = 'objective'
                if answer_text:
                    qobj['answer_text'] = answer_text
                if 'images' in locals() and images:
                    qobj['images'] = images
                    del images
                questions.append(qobj)
                i = j
            else:
                i += 1

    # Add diagnostic info to response to help debug parsing failures
    diagnostic = {
        'table_count': len(table_shapes),
        'table_shapes': table_shapes,
        'sample_table_texts': sample_table_texts,
        'paragraph_snippets': para_snippets,
        'parsed_questions': len(questions),
        'image_count': sum(len(q.get('images', [])) for q in questions)
    }
    # Summarize related parts (images, media) present in the docx package for debugging
    try:
        related = []
        for rid, part_obj in doc.part.related_parts.items():
            ct = getattr(part_obj, 'content_type', None) or str(type(part_obj))
            blob = getattr(part_obj, 'blob', None)
            size = len(blob) if blob else None
            related.append({'rel': rid, 'content_type': ct, 'size': size})
        diagnostic['related_parts'] = related
    except Exception:
        diagnostic['related_parts'] = []
    # Search all package parts for references to related part ids and image-like tags
    try:
        references = {rid: [] for rid, _ in doc.part.related_parts.items()}
        image_hint_parts = set()
        pkg = getattr(doc.part, 'package', None)
        if pkg is not None:
            for p in pkg.parts:
                try:
                    blob = getattr(p, 'blob', None)
                    if not blob:
                        continue
                    text = blob.decode('utf-8', errors='ignore')
                    for rid in list(references.keys()):
                        if rid in text:
                            references[rid].append(str(getattr(p, 'partname', p.partname)))
                    if any(k in text for k in ('<a:blip', 'pic:blipFill', '<v:imagedata', 'img src="')):
                        image_hint_parts.add(str(getattr(p, 'partname', p.partname)))
                except Exception:
                    continue
        diagnostic['references'] = references
        diagnostic['image_hint_parts'] = list(image_hint_parts)
    except Exception:
        diagnostic['references'] = {}
        diagnostic['image_hint_parts'] = []
    try:
        logger.info('scan-docx diagnostic: %s', diagnostic)
    except Exception:
        pass
    return {"questions": questions, "diagnostic": diagnostic}

@app.post("/api/template/generate-docx")
async def generate_docx(
    questions: list = Form(...),
    dept: str = Form(""),
    cc: str = Form(""),
    cn: str = Form(""),
    qpcode: str = Form(""),
    exam_title: str = Form("B.E., /B.Tech., DEGREE EXAMINATIONS, APRIL/MAY2024"),
    regulation: str = Form("Regulation 2023"),
    semester: str = Form("Second Semester"),
    excel_meta: str = Form(None),
    ocr_images: Optional[str] = Form(None),
    title_image_url: Optional[str] = Form(None),
    header_logo_url: Optional[str] = Form(None),
):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    import tempfile
    import os
    from io import BytesIO
    import base64, requests

    


    # If excel_meta is provided, parse and override cc/cn/dept/semester
    import json
    meta = {}
    if excel_meta:
        try:
            meta = json.loads(excel_meta)
        except Exception:
            meta = {}
    # Excel meta expected keys: course_code_name, department, semester (number or string)
    # Course code/name: prefer explicit keys, fall back to combined 'course_code_name' or UI params
    raw_code_name = meta.get('course_code_name')
    cc_from_excel = meta.get('course_code') or cc
    cn_from_excel = meta.get('course_name') or cn
    if not cn_from_excel and raw_code_name and isinstance(raw_code_name, str) and ' - ' in raw_code_name:
        parts = raw_code_name.split(' - ', 1)
        if not cc_from_excel:
            cc_from_excel = parts[0].strip()
        cn_from_excel = parts[1].strip()
    dept_from_excel = meta.get('department') or dept
    sem_from_excel = meta.get('semester') or semester
    # Convert semester number to words if needed
    def semester_to_words(sem):
        try:
            n = int(str(sem).strip())
            words = ["First", "Second", "Third", "Fourth", "Fifth", "Sixth", "Seventh", "Eighth"]
            if 1 <= n <= 8:
                return f"{words[n-1]} Semester"
        except Exception:
            pass
        return str(sem)
    sem_word = semester_to_words(sem_from_excel)

    doc = Document()

      # Insert banner image above semester line if provided
    logo_url = title_image_url if title_image_url else None
    if logo_url and isinstance(logo_url, str) and logo_url.strip():
        try:
            banner_tbl = doc.add_table(rows=1, cols=1)
            banner_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            banner_tbl.autofit = True
            banner_cell = banner_tbl.rows[0].cells[0]
            if logo_url.startswith('data:image/'):
                header, b64data = logo_url.split(',', 1)
                img_bytes = base64.b64decode(b64data)
                stream = BytesIO(img_bytes)
                banner_cell.paragraphs[0].add_run().add_picture(stream, width=Inches(6))
            elif logo_url.startswith('http://') or logo_url.startswith('https://'):
                resp = requests.get(logo_url, timeout=5)
                if resp.ok:
                    stream = BytesIO(resp.content)
                    banner_cell.paragraphs[0].add_run().add_picture(stream, width=Inches(6))
            banner_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            logger.exception("Failed to insert banner image; continuing without it")

    # Helpers
    def add_bold_line(text: str, align_center: bool = True, size: int = 12):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(size)
        if align_center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    def add_line(text: str, align_center: bool = True, size: int = 12, italic: bool = False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.italic = italic
        run.font.size = Pt(size)
        if align_center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p



    # If a compact header logo is provided, render exam title centered with logo at right
    if header_logo_url and isinstance(header_logo_url, str) and header_logo_url.strip():
        try:
            tbl = doc.add_table(rows=1, cols=2)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl.autofit = True
            # Title cell (left, centered)
            c_title = tbl.rows[0].cells[0]
            p_title = c_title.paragraphs[0]
            run_title = p_title.add_run((exam_title or "").strip() or "MODEL EXAM")
            run_title.bold = True
            run_title.font.size = Pt(20)
            run_title.underline = True
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Logo cell (right-aligned)
            c_logo = tbl.rows[0].cells[1]
            if header_logo_url.startswith('data:image/'):
                header, b64data = header_logo_url.split(',', 1)
                img_bytes = base64.b64decode(b64data)
                stream = BytesIO(img_bytes)
                c_logo.paragraphs[0].add_run().add_picture(stream, width=Inches(1.5))
            elif header_logo_url.startswith('http://') or header_logo_url.startswith('https://'):
                resp = requests.get(header_logo_url, timeout=5)
                if resp.ok:
                    stream = BytesIO(resp.content)
                    c_logo.paragraphs[0].add_run().add_picture(stream, width=Inches(1.5))
            c_logo.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        except Exception:
            logger.exception("Failed to insert header logo; falling back to text-only title")

    # (Banner already placed above; continue with course/meta lines)
    add_line(sem_word, True, 11, italic=True)
    add_line(dept_from_excel, True, 11, italic=True)
    # Display course code and name together when available
    display_course = cc_from_excel or ""
    if cn_from_excel:
        if display_course:
            display_course = f"{display_course} - {cn_from_excel}"
        else:
            display_course = cn_from_excel
    add_bold_line(display_course, True, 12)
    add_line(f"({regulation})", True, 11)

    # Time and Maximum Marks on same line, using a table for left/right alignment
    meta_tbl = doc.add_table(rows=1, cols=2)
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_tbl.autofit = True
    meta_tbl.allow_autofit = True
    # Left cell: Time
    p_time = meta_tbl.cell(0,0).paragraphs[0]
    run_time = p_time.add_run("Time: Three Hours")
    run_time.font.size = Pt(11)
    p_time.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Right cell: Maximum Marks
    p_marks = meta_tbl.cell(0,1).paragraphs[0]
    run_marks = p_marks.add_run("Maximum Marks: 100 Marks")
    run_marks.font.size = Pt(11)
    p_marks.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    try:
        meta_tbl.columns[0].width = Inches(3.2)
        meta_tbl.columns[1].width = Inches(3.6)
    except Exception:
        pass

   

    # PART-A title and marks in a single row using a table for alignment
    parta_tbl = doc.add_table(rows=1, cols=3)
    parta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    parta_tbl.autofit = True
    parta_tbl.allow_autofit = True
    # Left cell empty
    parta_tbl.cell(0,0).text = ""
    # Center cell: PART- A
    p_center = parta_tbl.cell(0,1).paragraphs[0]
    run_center = p_center.add_run("PART- A")
    run_center.bold = True
    run_center.font.size = Pt(12)
    p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Right cell: (10 x 2 = 20 Marks)
    p_right = parta_tbl.cell(0,2).paragraphs[0]
    run_right = p_right.add_run("(10 x 2 = 20 Marks)")
    run_right.bold = True
    run_right.font.size = Pt(12)
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Set column widths for better layout
    try:
        parta_tbl.columns[0].width = Inches(1.2)
        parta_tbl.columns[1].width = Inches(3.2)
        parta_tbl.columns[2].width = Inches(2.2)
    except Exception:
        pass

    table_a = doc.add_table(rows=1, cols=5)
    table_a.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_a.autofit = False
    hdr_cells = table_a.rows[0].cells
    hdr_cells[0].text = "Q.No."
    hdr_cells[1].text = "Answer ALL Questions"
    hdr_cells[2].text = "CO"
    hdr_cells[3].text = "BTL"
    hdr_cells[4].text = "Marks"
    # Column widths
    widths_a = [Inches(0.5), Inches(5), Inches(0.5), Inches(0.6), Inches(0.6)]
    for i, w in enumerate(widths_a):
        for row in table_a.rows:
            row.cells[i].width = w
    # Align header columns: Q.No., CO, BTL, Marks centered; Answer ALL left
    hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[1].paragraphs[0].paragraph_format.left_indent = Inches(1)
    hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for c in hdr_cells:
        for r in c.paragraphs[0].runs:
            r.bold = True

    # Normalize incoming questions to dicts
    import json, random
    def _normalize(raw):
        out = []
        def _push(item):
            if isinstance(item, str):
                try:
                    item = json.loads(item)
                except Exception:
                    item = {"text": item}
            if isinstance(item, list):
                for sub in item:
                    _push(sub)
            elif isinstance(item, dict):
                out.append(item)
        # entry point
        if isinstance(raw, str) or isinstance(raw, dict) or isinstance(raw, list):
            _push(raw)
        else:
            try:
                _push(json.loads(str(raw)))
            except Exception:
                pass
        return out

    _questions = _normalize(questions)
    # Parse optional OCR images map: { index_or_id: dataUrl }
    import json as _json
    ocr_map = {}
    try:
        if ocr_images:
            ocr_map = _json.loads(ocr_images)
            if not isinstance(ocr_map, dict):
                ocr_map = {}
    except Exception:
        logger.exception("Failed to parse ocr_images payload")
    logger.info("generate_docx called: received %d question(s)", len(_questions))
    # Normalize common image keys and log per-question image presence (without dumping full base64)
    try:
        image_count = 0
        for i, q in enumerate(_questions):
            if not isinstance(q, dict):
                logger.debug("Question %s is not a dict (type=%s)", i, type(q))
                continue
            # Support alternate keys the frontend might send
            for alt in ('image', 'img', 'imageUrl', 'img_url'):
                if alt in q and 'image_url' not in q:
                    q['image_url'] = q.get(alt)
                    logger.debug("Normalized image key '%s' -> 'image_url' for question %s", alt, i)

            img = q.get('image_url')
            if img:
                image_count += 1
                try:
                    kind = 'data' if isinstance(img, str) and img.startswith('data:') else ('http' if isinstance(img, str) and img.startswith('http') else type(img))
                    preview = (img[:80] + '...') if isinstance(img, str) and len(img) > 80 else str(img)
                    logger.info("Question %s: has image_url (kind=%s, preview=%s)", i, kind, preview)
                except Exception:
                    logger.exception("Error while logging image preview for question %s", i)
            else:
                logger.debug("Question %s: no image_url", i)

        logger.info("Found %d questions with image_url", image_count)
    except Exception:
        logger.exception("Error while counting/normalizing image URLs")

    # Choose up to 10 random Part-A questions; fall back sensibly
    a_questions = [q for q in _questions if str(q.get('part', '')).upper() == 'A']
    if not a_questions:
        # heuristic fallback by marks/type if part missing
        a_questions = [
            q for q in _questions
            if str(q.get('marks', '')).strip().lower() in ("2", "2m", "2 marks", "2 mark", "2-marks")
            or str(q.get('type', '')).strip().lower() in ("a", "part-a", "short", "objective", "two", "2")
        ]
    if not a_questions:
        a_questions = _questions

    random.shuffle(a_questions)
    selected_a = a_questions[:10]

    # Helpers to read common field names
    def _first_non_empty(d: dict, keys: list[str]):
        for k in keys:
            if k in d and d[k] is not None and str(d[k]).strip() != "":
                return str(d[k])
        return ""

    # Separate descriptive and objective questions for PART-A
    desc_qs = [q for q in _questions if isinstance(q, dict) and (
        str(q.get('type', '')).lower() in ('descriptive', 'd', 'desc', 'long', 'theory') or
        (str(q.get('part', '')).upper() == 'A' and str(q.get('marks', '2')).strip() != '2')
    )]
    obj_qs = [q for q in _questions if isinstance(q, dict) and (
        str(q.get('type', '')).lower() in ('objective', 'o', 'obj', 'short', 'mcq', 'one', '2', 'two') or
        (str(q.get('part', '')).upper() == 'A' and str(q.get('marks', '2')).strip() == '2')
    )]
    # Fallbacks if not enough
    if not desc_qs:
        desc_qs = [q for q in _questions if isinstance(q, dict)]
    if not obj_qs:
        obj_qs = [q for q in _questions if isinstance(q, dict)]
    random.shuffle(desc_qs)
    random.shuffle(obj_qs)

    # Prepare BTL shared value for questions 5-10 (must be 3,4 or 5)
    btl_shared = random.choice([3, 4, 5])
    idx = 1
    import re

    # Use the first 10 questions as received (frontend order)
    import base64, requests
    from io import BytesIO
    # OCR helper
    def _ocr_data_url(data_url: str) -> Optional[str]:
        try:
            import pytesseract
            from PIL import Image
            header, b64data = data_url.split(',', 1)
            img_bytes = base64.b64decode(b64data)
            with BytesIO(img_bytes) as bio:
                img = Image.open(bio)
                # Convert to RGB to avoid issues
                if img.mode not in ("RGB", "L"):
                    img = img.convert("RGB")
                text = pytesseract.image_to_string(img)
                cleaned = text.strip()
                return cleaned if cleaned else None
        except Exception:
            logger.exception("OCR failed for provided data URL")
            return None
    for i, q in enumerate(_questions[:10]):
        row_cells = table_a.add_row().cells
        for j, w in enumerate(widths_a):
            row_cells[j].width = w

        # Question number
        row_cells[0].text = str(idx)

        # Question text without any label prefix
        text = _first_non_empty(q, ['text', 'question_text', 'question', 'q', 'title', 'body', 'content'])
        if text:
            text = re.sub(r'^\s*[DO]\.[\s-]*', '', text, flags=re.IGNORECASE)

        # Add text to cell, then image if present
        p = row_cells[1].paragraphs[0]
        if text:
            p.add_run(text)
        img_url = q.get('image_url')
        if img_url:
            try:
                # If OCR is requested and we have a matching data URL, try OCR
                do_ocr = bool(q.get('image_ocr'))
                if do_ocr and isinstance(img_url, str) and img_url.startswith('data:image/'):
                    ocr_text = _ocr_data_url(img_url)
                    if ocr_text:
                        p.add_run("\n" + ocr_text)
                        logger.info("Inserted OCR text for question index %s", idx)
                        raise StopIteration  # Skip image insertion below
                if img_url.startswith('data:image/'):
                    # data URL
                    header, b64data = img_url.split(',', 1)
                    img_bytes = base64.b64decode(b64data)
                    ext = '.png' if 'png' in header else '.jpg'
                    img_stream = BytesIO(img_bytes)
                    p.add_run().add_picture(img_stream, width=Inches(2.5))
                    logger.info("Inserted data:image for question index %s (ext=%s)", idx, ext)
                elif img_url.startswith('http'):
                    resp = requests.get(img_url)
                    if resp.ok:
                        content_type = resp.headers.get('content-type','')
                        ext = '.png' if 'png' in content_type else '.jpg'
                        img_stream = BytesIO(resp.content)
                        p.add_run().add_picture(img_stream, width=Inches(2.5))
                        logger.info("Fetched and inserted remote image for question index %s (content-type=%s)", idx, content_type)
            except StopIteration:
                pass
            except Exception as e:
                logger.exception("Failed to insert image for question index %s, img_url=%s", idx, img_url)
                p.add_run(" [Image error]")

        # CO: use from question if present, else fallback to mapping
        co_val = q.get('co') or q.get('CO') or q.get('course_outcome')
        if co_val:
            row_cells[2].text = str(co_val)
        else:
            co_num = (i // 2) + 1
            row_cells[2].text = f"CO{co_num}"

        # BTL logic: random for first 4 questions, shared value for 5-10
        if i < 4:
            btl_val = q.get('btl') or q.get('BTL') or random.choice([1, 2, 3, 4, 5])
        else:
            btl_val = q.get('btl') or q.get('BTL') or btl_shared
        # Remove 'BTL' prefix if present, only show number
        btl_str = str(btl_val)
        if btl_str.upper().startswith('BTL'):
            btl_str = btl_str[3:]
        row_cells[3].text = btl_str.strip()

        # Marks: always 2 for Part-A
        row_cells[4].text = "2"

        # center-align numeric/small columns
        for p in row_cells[0].paragraphs + row_cells[2].paragraphs + row_cells[3].paragraphs + row_cells[4].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        idx += 1
    # PART-B title and marks in a single row using a table for alignment
    partb_tbl = doc.add_table(rows=1, cols=3)
    partb_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    partb_tbl.autofit = True
    partb_tbl.allow_autofit = True
    partb_tbl.cell(0,0).text = ""
    p_center_b = partb_tbl.cell(0,1).paragraphs[0]
    run_center_b = p_center_b.add_run("PART- B")
    run_center_b.bold = True
    run_center_b.font.size = Pt(12)
    p_center_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_right_b = partb_tbl.cell(0,2).paragraphs[0]
    run_right_b = p_right_b.add_run("(5 x 16 = 80 Marks)")
    run_right_b.bold = True
    run_right_b.font.size = Pt(12)
    p_right_b.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    try:
        partb_tbl.columns[0].width = Inches(1.2)
        partb_tbl.columns[1].width = Inches(3.2)
        partb_tbl.columns[2].width = Inches(2.2)
    except Exception:
        pass

    table_b = doc.add_table(rows=1, cols=5)
    table_b.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_b.autofit = False
    table_b.left_indent = Inches(0.1)
    hdr_cells = table_b.rows[0].cells
    hdr_cells[0].text = "Q.No."
    hdr_cells[1].text = "Answer All Questions"
    hdr_cells[2].text = "CO"
    hdr_cells[3].text = "BTL"
    hdr_cells[4].text = "Marks"
    # Apply same header widths/alignment as Part A
    widths_b = [Inches(0.5), Inches(5), Inches(0.5), Inches(0.6), Inches(0.6)]
    for i, w in enumerate(widths_b):
        for row in table_b.rows:
            row.cells[i].width = w
    hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[1].paragraphs[0].paragraph_format.left_indent = Inches(1)
    hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for c in hdr_cells:
        for r in c.paragraphs[0].runs:
            r.bold = True
    # Render Part B with (a), OR row, (b) for each pair
    b_pairs = []
    temp_pair = []
    for q in _questions:
        if isinstance(q, dict) and q.get('part', '').upper() == 'B':
            temp_pair.append(q)
            if len(temp_pair) == 2:
                b_pairs.append(temp_pair)
                temp_pair = []
    # Always print 5 pairs: 11, 12, 13, 14, 15, even if some pairs are missing (insert empty rows)
    for pair_idx in range(5):
        base_no = str(11 + pair_idx)
        pair = b_pairs[pair_idx] if pair_idx < len(b_pairs) else None
        if pair:
            try:
                pair_sorted = sorted(pair, key=lambda q: str(q.get('sub','')))
            except Exception:
                pair_sorted = pair
            qa, qb = pair_sorted
        else:
            qa, qb = None, None
        # (a) row
        row_a = table_b.add_row().cells
        for i, w in enumerate(widths_b):
            row_a[i].width = w
        row_a[0].text = f"{base_no} a"
        p_a = row_a[1].paragraphs[0]
        p_a.paragraph_format.left_indent = Inches(0.0)
        if qa:
            # Add question text
            p_a.add_run(_first_non_empty(qa, ['text','question_text','question','q','title','body','content']))
            # Add image if present
            img_url = qa.get('image_url')
            if img_url:
                try:
                    do_ocr = bool(qa.get('image_ocr'))
                    if do_ocr and isinstance(img_url, str) and img_url.startswith('data:image/'):
                        ocr_text = _ocr_data_url(img_url)
                        if ocr_text:
                            p_a.add_run("\n" + ocr_text)
                            logger.info("Inserted OCR text for PART-B (a) %s", base_no)
                            raise StopIteration
                    if img_url.startswith('data:image/'):
                        header, b64data = img_url.split(',', 1)
                        img_bytes = base64.b64decode(b64data)
                        ext = '.png' if 'png' in header else '.jpg'
                        img_stream = BytesIO(img_bytes)
                        p_a.add_run().add_picture(img_stream, width=Inches(2.5))
                        logger.info("Inserted data:image for PART-A question %s (ext=%s)", idx, ext)
                    elif img_url.startswith('http'):
                        resp = requests.get(img_url)
                        if resp.ok:
                            content_type = resp.headers.get('content-type','')
                            ext = '.png' if 'png' in content_type else '.jpg'
                            img_stream = BytesIO(resp.content)
                            p_a.add_run().add_picture(img_stream, width=Inches(2.5))
                            logger.info("Fetched and inserted PART-A remote image for %s (content-type=%s)", idx, content_type)
                except StopIteration:
                    pass
                except Exception as e:
                    logger.exception("Failed to insert PART-A image for %s, img_url=%s", idx, img_url)
                    p_a.add_run(" [Image error]")
            row_a[2].text = _first_non_empty(qa, ['co','course_outcomes','courseOutcome','course_outcome','co_code'])
            row_a[3].text = _first_non_empty(qa, ['btl','bloom','bloom_level','bt','bt_level'])
            row_a[4].text = _first_non_empty(qa, ['marks','mark','score','points'])
        else:
            row_a[2].text = row_a[3].text = row_a[4].text = ""
        for p in (
            row_a[0].paragraphs + row_a[2].paragraphs + row_a[3].paragraphs + row_a[4].paragraphs
        ):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # OR row (spanning all columns, centered)
        or_row = table_b.add_row().cells
        or_row[0].merge(or_row[-1])
        p_or = or_row[0].paragraphs[0]
        p_or.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_or.add_run("(OR)").bold = True

        # (b) row
        row_b = table_b.add_row().cells
        for i, w in enumerate(widths_b):
            row_b[i].width = w
        row_b[0].text = f"{base_no} b"
        p_b = row_b[1].paragraphs[0]
        p_b.paragraph_format.left_indent = Inches(0.0)
        if qb:
            # Add question text
            p_b.add_run(_first_non_empty(qb, ['text','question_text','question','q','title','body','content']))
            # Add image if present
            img_url = qb.get('image_url')
            if img_url:
                try:
                    do_ocr = bool(qb.get('image_ocr'))
                    if do_ocr and isinstance(img_url, str) and img_url.startswith('data:image/'):
                        ocr_text = _ocr_data_url(img_url)
                        if ocr_text:
                            p_b.add_run("\n" + ocr_text)
                            logger.info("Inserted OCR text for PART-B (b) %s", base_no)
                            raise StopIteration
                    if img_url.startswith('data:image/'):
                        header, b64data = img_url.split(',', 1)
                        img_bytes = base64.b64decode(b64data)
                        ext = '.png' if 'png' in header else '.jpg'
                        img_stream = BytesIO(img_bytes)
                        p_b.add_run().add_picture(img_stream, width=Inches(2.5))
                        logger.info("Inserted data:image for PART-B question %s (ext=%s)", idx, ext)
                    elif img_url.startswith('http'):
                        resp = requests.get(img_url)
                        if resp.ok:
                            content_type = resp.headers.get('content-type','')
                            ext = '.png' if 'png' in content_type else '.jpg'
                            img_stream = BytesIO(resp.content)
                            p_b.add_run().add_picture(img_stream, width=Inches(2.5))
                            logger.info("Fetched and inserted PART-B remote image for %s (content-type=%s)", idx, content_type)
                except StopIteration:
                    pass
                except Exception as e:
                    logger.exception("Failed to insert PART-B image for %s, img_url=%s", idx, img_url)
                    p_b.add_run(" [Image error]")
            row_b[2].text = _first_non_empty(qb, ['co','course_outcomes','courseOutcome','course_outcome','co_code'])
            row_b[3].text = _first_non_empty(qb, ['btl','bloom','bloom_level','bt','bt_level'])
            row_b[4].text = _first_non_empty(qb, ['marks','mark','score','points'])
        else:
            row_b[2].text = row_b[3].text = row_b[4].text = ""
        for p in (
            row_b[0].paragraphs + row_b[2].paragraphs + row_b[3].paragraphs + row_b[4].paragraphs
        ):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # PART-C (optional single question worth 10 marks)
    c_questions = [q for q in _questions if isinstance(q, dict) and (
        str(q.get('part','')).upper() == 'C' or
        str(q.get('baseNumber','')).strip() == '16' or
        str(q.get('number','')).strip().startswith('16')
    )]
    if c_questions:
        partc_tbl = doc.add_table(rows=1, cols=3)
        partc_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        partc_tbl.autofit = True
        partc_tbl.allow_autofit = True
        partc_tbl.cell(0,0).text = ""
        p_center_c = partc_tbl.cell(0,1).paragraphs[0]
        run_center_c = p_center_c.add_run("PART- C")
        run_center_c.bold = True
        run_center_c.font.size = Pt(12)
        p_center_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_right_c = partc_tbl.cell(0,2).paragraphs[0]
        run_right_c = p_right_c.add_run("    (1 x 10 = 10 Marks)")
        run_right_c.bold = True
        run_right_c.font.size = Pt(12)
        p_right_c.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        try:
            partc_tbl.columns[0].width = Inches(1.2)
            partc_tbl.columns[1].width = Inches(3.2)
            partc_tbl.columns[2].width = Inches(2.2)
        except Exception:
            pass

        table_c = doc.add_table(rows=1, cols=5)
        table_c.alignment = WD_TABLE_ALIGNMENT.CENTER
        table_c.autofit = False
        hdr_cells = table_c.rows[0].cells
        hdr_cells[0].text = "Q.No."
        hdr_cells[1].text = "Answer All Questions"
        hdr_cells[2].text = "CO"
        hdr_cells[3].text = "BTL"
        hdr_cells[4].text = "Marks"
        # Apply same header widths/alignment as Part A
        widths_c = [Inches(0.5), Inches(5), Inches(0.5), Inches(0.6), Inches(0.6)]
        for i, w in enumerate(widths_c):
            for row in table_c.rows:
                row.cells[i].width = w
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[1].paragraphs[0].paragraph_format.left_indent = Inches(1)
        hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for c in hdr_cells:
            for r in c.paragraphs[0].runs:
                r.bold = True
        # render Part C as an (a) / (OR) / (b) block (lowercase a/b)
        qa = None
        qb = None
        if len(c_questions) >= 2:
            try:
                pair_sorted = sorted(c_questions, key=lambda q: str(q.get('sub','')))
            except Exception:
                pair_sorted = c_questions[:2]
            qa, qb = pair_sorted[0], pair_sorted[1]
        else:
            qa = c_questions[0]

        base_no = _first_non_empty(qa if qa else (qb or {}), ['baseNumber','number']) or '16'

        # (a) row
        row_a = table_c.add_row().cells
        for i, w in enumerate(widths_c):
            row_a[i].width = w
        row_a[0].text = f"{base_no} a"
        p_a = row_a[1].paragraphs[0]
        p_a.paragraph_format.left_indent = Inches(0.0)
        if qa:
            p_a.add_run(_first_non_empty(qa, ['text','question_text','question','q','title','body','content']))
            img_url = qa.get('image_url')
            if img_url:
                try:
                    do_ocr = bool(qa.get('image_ocr'))
                    if do_ocr and isinstance(img_url, str) and img_url.startswith('data:image/'):
                        ocr_text = _ocr_data_url(img_url)
                        if ocr_text:
                            p_a.add_run("\n" + ocr_text)
                            logger.info("Inserted OCR text for PART-C (a) %s", base_no)
                            raise StopIteration
                    if img_url.startswith('data:image/'):
                        header, b64data = img_url.split(',', 1)
                        img_bytes = base64.b64decode(b64data)
                        img_stream = BytesIO(img_bytes)
                        p_a.add_run().add_picture(img_stream, width=Inches(2.5))
                    elif img_url.startswith('http'):
                        resp = requests.get(img_url)
                        if resp.ok:
                            img_stream = BytesIO(resp.content)
                            p_a.add_run().add_picture(img_stream, width=Inches(2.5))
                except StopIteration:
                    pass
                except Exception:
                    logger.exception("Failed to insert PART-C (a) image for %s, img_url=%s", base_no, img_url)
                    p_a.add_run(" [Image error]")
        row_a[2].text = _first_non_empty(qa or {}, ['co','course_outcomes','courseOutcome','course_outcome','co_code'])
        row_a[3].text = _first_non_empty(qa or {}, ['btl','bloom','bloom_level','bt','bt_level'])
        row_a[4].text = _first_non_empty(qa or {}, ['marks','mark','score','points']) or '10'
        for p in (row_a[0].paragraphs + row_a[2].paragraphs + row_a[3].paragraphs + row_a[4].paragraphs):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # OR row (spanning all columns, centered)
        or_row = table_c.add_row().cells
        or_row[0].merge(or_row[-1])
        p_or = or_row[0].paragraphs[0]
        p_or.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_or.add_run("(OR)").bold = True

        # (b) row
        row_b = table_c.add_row().cells
        for i, w in enumerate(widths_c):
            row_b[i].width = w
        row_b[0].text = f"{base_no} b"
        p_b = row_b[1].paragraphs[0]
        p_b.paragraph_format.left_indent = Inches(0.0)
        if qb:
            p_b.add_run(_first_non_empty(qb, ['text','question_text','question','q','title','body','content']))
            img_url = qb.get('image_url')
            if img_url:
                try:
                    do_ocr = bool(qb.get('image_ocr'))
                    if do_ocr and isinstance(img_url, str) and img_url.startswith('data:image/'):
                        ocr_text = _ocr_data_url(img_url)
                        if ocr_text:
                            p_b.add_run("\n" + ocr_text)
                            logger.info("Inserted OCR text for PART-C (b) %s", base_no)
                            raise StopIteration
                    if img_url.startswith('data:image/'):
                        header, b64data = img_url.split(',', 1)
                        img_bytes = base64.b64decode(b64data)
                        img_stream = BytesIO(img_bytes)
                        p_b.add_run().add_picture(img_stream, width=Inches(2.5))
                    elif img_url.startswith('http'):
                        resp = requests.get(img_url)
                        if resp.ok:
                            img_stream = BytesIO(resp.content)
                            p_b.add_run().add_picture(img_stream, width=Inches(2.5))
                except StopIteration:
                    pass
                except Exception:
                    logger.exception("Failed to insert PART-C (b) image for %s, img_url=%s", base_no, img_url)
                    p_b.add_run(" [Image error]")
        row_b[2].text = _first_non_empty(qb or {}, ['co','course_outcomes','courseOutcome','course_outcome','co_code'])
        row_b[3].text = _first_non_empty(qb or {}, ['btl','bloom','bloom_level','bt','bt_level'])
        row_b[4].text = _first_non_empty(qb or {}, ['marks','mark','score','points']) or '10'
        for p in (row_b[0].paragraphs + row_b[2].paragraphs + row_b[3].paragraphs + row_b[4].paragraphs):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(" ")
    doc.add_paragraph("******************").bold = True
    # Footer
    doc.add_paragraph(f"  {qpcode}").bold = True
    # Save to temp file
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name
    return FileResponse(tmp_path, filename="question_paper.docx")

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=4000)
